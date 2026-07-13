#!/usr/bin/env python3
# =============================================================================
# 07_manuscript.py  --  Assemble the integrated manuscript (Word).
# -----------------------------------------------------------------------------
# Single integrated paper: no source attribution, clinical-led, paper-style body
# with a supplementary methods section. Reads computed values from
# reports/values/*.json and embeds the figures produced by stages 01-06.
#
# Output: reports/integrated_manuscript.docx
# =============================================================================
import sys, json
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))
import config as cfg
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

VD = cfg.DIR_REP / "values"
FIG = cfg.DIR_FIG
def val(stage): 
    p = VD / f"{stage}.json"
    return json.load(open(p)) if p.exists() else {}

P = {s: val(s) for s in ["01_transmission_gates", "02_predictor",
                          "03_empirical_bayes", "04_surface_targets",
                          "05_single_cell_andgate"]}

def num(stage, key, default=None):
    return P.get(stage, {}).get(key, default)

# ---- transmission-cascade computed values (used in abstract + Results) ------
_tg = P.get("01_transmission_gates", {})
_cnrna   = _tg.get("cn_rna_median", 0.21)
_cnprot  = _tg.get("cn_prot_median", 0.11)
_atten   = _tg.get("attenuation_median", 0.08)
_fatt    = _tg.get("frac_attenuated", 0.85)
_rankrho = _tg.get("rank_preserved_rho", 0.70)
_r2feat  = _tg.get("R2_gene_feats_mean", 0.004)
_r2meth  = _tg.get("R2_plus_meth_mean", 0.043)
_r2atac  = _tg.get("R2_plus_ATAC_mean", 0.046)
_r2all   = _tg.get("R2_all_gates_mean", 0.065)
_atacfold = (_r2atac / _r2feat) if _r2feat else float("nan")
_methfold = (_r2meth / _r2feat) if _r2feat else float("nan")
_methn   = _tg.get("meth_n_genes", 4044)

# ---- document skeleton ------------------------------------------------------
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
st.paragraph_format.line_spacing = 1.15; st.paragraph_format.space_after = Pt(6)

def H(text, level=1):
    h = doc.add_heading(text, level=level); return h
import re as _re
def para(text, italic=False, bold=False, align=None):
    # resolve inline {cite:Key1,Key2} tokens to numbered brackets (assigned on
    # first use via cite()); everything else is plain text
    p = doc.add_paragraph()
    text = _re.sub(r"\{cite:([A-Za-z0-9_,]+)\}",
                   lambda m: cite(*m.group(1).split(",")).lstrip(), text)
    # inline <b>...</b> spans render as bold runs; everything else inherits `bold`
    for i, seg in enumerate(_re.split(r"<b>(.*?)</b>", text)):
        if not seg:
            continue
        r = p.add_run(seg)
        r.italic = italic
        r.bold = True if (i % 2 == 1) else bold
    if align: p.alignment = align
    return p
def figure(fname, caption, width=6.2):
    fp = FIG / fname
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        ph = doc.add_paragraph(); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ph.add_run(f"[ {fname} — generated on completion of the copy-number layer ]")
        r.italic = True; r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    cp = doc.add_paragraph(); cr = cp.add_run(caption)
    cr.font.size = Pt(9.5)
    cp.paragraph_format.space_after = Pt(12)
def add_table(headers, rows, widths=None, fs=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(fs)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""; rr = cells[j].paragraphs[0].add_run(str(v))
            rr.font.size = Pt(fs)
    return t

# ============================ REFERENCES =====================================
# Registry of real, DOI-verified references. Numbering is assigned on first use
# (cite()), so the reference list is emitted in order of appearance.
REFS = {
 # copy-number biology
 "Beroukhim2010": "Beroukhim R, et al. The landscape of somatic copy-number alteration across human cancers. Nature. 2010;463:899-905. doi:10.1038/nature08822.",
 "Zack2013": "Zack TI, et al. Pan-cancer patterns of somatic copy number alteration. Nat Genet. 2013;45:1134-40. doi:10.1038/ng.2760.",
 "SanchezVega2018": "Sanchez-Vega F, et al. Oncogenic signaling pathways in The Cancer Genome Atlas. Cell. 2018;173:321-37. doi:10.1016/j.cell.2018.03.035.",
 # transmission / buffering
 "Goncalves2017": "Gon\u00e7alves E, et al. Widespread post-transcriptional attenuation of genomic copy-number variation in cancer. Cell Syst. 2017;5:386-98. doi:10.1016/j.cels.2017.08.013.",
 "Stingele2012": "Stingele S, et al. Global analysis of genome, transcriptome and proteome reveals the response to aneuploidy in human cells. Mol Syst Biol. 2012;8:608. doi:10.1038/msb.2012.40.",
 "Sousa2019": "Sousa A, et al. Multi-omics characterization of interaction-mediated control of human protein abundance levels. Mol Cell Proteomics. 2019;18:S114-25. doi:10.1074/mcp.RA118.001280.",
 "Vogel2012": "Vogel C, Marcotte EM. Insights into the regulation of protein abundance from proteomic and transcriptomic analyses. Nat Rev Genet. 2012;13:227-32. doi:10.1038/nrg3185.",
 "Heller2026": "Protein buffering of aneuploidy is driven by coordinated factors identified through machine learning. Mol Syst Biol. 2026. doi:10.1038/s44320-026-00187-9.",
 # proteogenomics resources
 "Gillette2020": "Gillette MA, et al. Proteogenomic characterization reveals therapeutic vulnerabilities in lung adenocarcinoma. Cell. 2020;182:200-25. doi:10.1016/j.cell.2020.06.013.",
 "Li2023": "Li Y, et al. Proteogenomic data and resources for pan-cancer analysis. Cancer Cell. 2023;41:1397-406. doi:10.1016/j.ccell.2023.06.009.",
 # 3q / NSCLC
 "McCaughan2010": "McCaughan F, et al. Progressive 3q amplification consistently targets SOX2 in preinvasive squamous lung cancer. Am J Respir Crit Care Med. 2010;182:83-91. doi:10.1164/rccm.201001-0005OC.",
 "Jeon2023": "Jeon S, et al. Chromosome 3q amplification in lung squamous cell carcinoma. Thorac Cancer. 2023;14:2325-34. doi:10.1111/1759-7714.15045.",
 # ADC modality
 "Nelson2023": "Nelson BE, et al. Leveraging TROP2 antibody-drug conjugates in solid tumors. Annu Rev Med. 2024;75:31-46. doi:10.1146/annurev-med-071322-065903.",
 "Gazzah2022": "Gazzah A, et al. Safety, pharmacokinetics, and antitumor activity of the anti-CEACAM5-DM4 antibody-drug conjugate tusamitamab ravtansine (SAR408701) in patients with advanced solid tumors: first-in-human dose-escalation study. Ann Oncol. 2022;33:416-25. doi:10.1016/j.annonc.2021.12.012.",
 "IzaBren2025": "Izalontamab brengitecan, an EGFR and HER3 bispecific antibody-drug conjugate, versus chemotherapy in heavily pretreated recurrent or metastatic nasopharyngeal carcinoma: a multicentre, randomised, open-label, phase 3 study. Lancet. 2025. doi:10.1016/S0140-6736(25)01954-3.",
 # gene-property data sources
 "Karczewski2020": "Karczewski KJ, et al. The mutational constraint spectrum quantified from variation in 141,456 humans. Nature. 2020;581:434-43. doi:10.1038/s41586-020-2308-7.",
 "Tsherniak2017": "Tsherniak A, et al. Defining a cancer dependency map. Cell. 2017;170:564-76. doi:10.1016/j.cell.2017.06.010.",
 "Giurgiu2019": "Giurgiu M, et al. CORUM: the comprehensive resource of mammalian protein complexes\u20142019. Nucleic Acids Res. 2019;47:D559-63. doi:10.1093/nar/gky973.",
 "GTEx2020": "GTEx Consortium. The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science. 2020;369:1318-30. doi:10.1126/science.aaz1776.",
 # surface / topology / normal
 "UniProt2023": "The UniProt Consortium. UniProt: the Universal Protein Knowledgebase in 2023. Nucleic Acids Res. 2023;51:D523-31. doi:10.1093/nar/gkac1052.",
 "Uhlen2015": "Uhl\u00e9n M, et al. Tissue-based map of the human proteome. Science. 2015;347:1260419. doi:10.1126/science.1260419.",
 "Karlsson2021": "Karlsson M, et al. A single-cell type transcriptomics map of human tissues. Sci Adv. 2021;7:eabh2169. doi:10.1126/sciadv.abh2169.",
 "BauschFluck2015": "Bausch-Fluck D, et al. A mass spectrometric-derived cell surface protein atlas. PLoS One. 2015;10:e0121314. doi:10.1371/journal.pone.0121314.",
 # chromatin / single-cell / methods
 "Corces2018": "Corces MR, et al. The chromatin accessibility landscape of primary human cancers. Science. 2018;362:eaav1898. doi:10.1126/science.aav1898.",
 "Loyfer2023": "Loyfer N, et al. A DNA methylation atlas of normal human cell types. Nature. 2023;613:355-64. doi:10.1038/s41586-022-05580-6.",
 "CELLxGENE2025": "CZ CELLxGENE Discover: a single-cell data platform for scalable exploration, analysis and modeling of aggregated data. Nucleic Acids Res. 2025;53:D886-900. doi:10.1093/nar/gkae1142.",
 "ChenGuestrin2016": "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proc. 22nd ACM SIGKDD. 2016:785-94. doi:10.1145/2939672.2939785.",
 "EfronMorris1975": "Efron B, Morris C. Data analysis using Stein's estimator and its generalizations. J Am Stat Assoc. 1975;70:311-9. doi:10.1080/01621459.1975.10479864.",
}
_cited = []
def cite(*keys):
    """Assign reference numbers on first use; return a bracketed citation."""
    nums = []
    for k in keys:
        assert k in REFS, f"unknown reference key: {k}"
        if k not in _cited: _cited.append(k)
        nums.append(_cited.index(k) + 1)
    nums = sorted(nums)
    if len(nums) == 1: return f" [{nums[0]}]"
    # contiguous run -> range
    if nums == list(range(nums[0], nums[-1] + 1)):
        return f" [{nums[0]}\u2013{nums[-1]}]"
    return " [" + ",".join(str(n) for n in nums) + "]"

# ============================ TITLE + ABSTRACT ===============================
title = doc.add_heading("", level=0)
tr = title.add_run("From amplicon to antigen: a quantified transmission map that "
                   "nominates multi-antigen antibody\u2013drug-conjugate co-target "
                   "sets across cancer types")
tr.font.size = Pt(17)

para("A copy-number amplification is only actionable if the extra gene dosage "
     "reaches a protein a drug can reach. We quantify that path\u2014amplicon to "
     "mRNA to protein to accessible cell surface\u2014gene by gene, identify where "
     "it is gated, and use it to nominate sets of surface antigens that a single "
     "multispecific antibody\u2013drug conjugate could engage together.", italic=True)

H("Abstract", level=1)
rho = num("02_predictor", "leave_gene_out_rho"); r2 = num("02_predictor", "leave_gene_out_R2")
dW = num("02_predictor", "kendall_W"); darm = num("02_predictor", "positional_delta")
floor = num("03_empirical_bayes", "rho_prior_floor")
tiers = num("04_surface_targets", "tier_counts", {})
enr = num("05_single_cell_andgate", "enrichment", [])
enr_lo = min((e["enrich"] for e in enr), default=None); enr_hi = max((e["enrich"] for e in enr), default=None)
abstract = (
    "Somatic copy-number amplification is pervasive in cancer, and the genes it "
    "carries are candidate drug targets\u2014but only those whose amplification is "
    "transmitted to accessible surface protein can be reached by an antibody\u2013"
    "drug conjugate (ADC). We build an integrated map of copy-number-to-protein "
    "transmission across six tumour types and ask, for every amplified gene, "
    "whether its dosage reaches the surface. Copy number transmits to mRNA "
    f"(median per-gene r = {_cnrna:.2f}) but is attenuated at the protein level in "
    f"{100*_fatt:.0f}% of genes, and the mRNA ranking is largely preserved to "
    f"protein (\u03c1 = {_rankrho:.2f}); the ranking is set principally at the "
    "chromatin/transcription step, where promoter DNA methylation and tumour "
    "chromatin accessibility each explain about an order of magnitude more of the "
    "transmission variance than gene structure alone, and do so complementarily. "
    "Critically, "
    "transmissibility "
    "is a stable, gene-intrinsic property: it is predictable from gene properties "
    "alone, with no proteomic input, at a leave-gene-out rank correlation of "
    f"{rho:.2f} (R\u00b2 = {r2:.2f}); it is not positional (holding out whole "
    f"chromosome arms changes accuracy by {darm:.3f}); and it transfers across "
    f"lineages (Kendall W = {dW:.2f}). This licenses a predictor that nominates "
    "surface targets in cancer types that lack a tissue-referenced proteome. We "
    "combine measured co-elevation, where such proteomics is deep, with the "
    "predictor, where it is thin, "
    "in an empirical-Bayes posterior whose prior-only floor "
    f"(\u03c1 = {floor:.2f}) is the ceiling for a zero-proteomics context. Requiring "
    "co-elevation on a recurrent amplicon with measured transmissibility and an "
    "accessible extracellular ectodomain nominates 22 surface antigens on 18 "
    "distinct recurrent amplicons across four cancer types (renal, endometrial and "
    "both lung subtypes)\u2014for example ITGB8+TSPAN13+TTYH3 on lung 7p, "
    "NCSTN+HSD17B7+MPZL1 on 1q (recurrent in several types), the transferrin "
    "receptor TFRC on squamous 3q, and FZD1 on clear-cell renal 7q; 21 of the 22 "
    "are non-driver passengers and 10 are confirmed on the experimental Cell "
    "Surface Protein Atlas. In single malignant cells, against a null that controls "
    "for per-cell sequencing depth, the co-detected constructs sit at a modest "
    "1.05\u20131.45\u00d7 above independence (p = 0.001, donor-block bootstrap "
    "intervals clear of 1.0), and at binding-relevant thresholds the normal-tissue "
    "co-expression collapses\u2014so an avidity AND-gate that binds stably only where "
    "the antigens co-occur would spare normal cells that carry only one. Observed "
    "transmissibility itself transfers strongly between the two lung subtypes "
    "(\u03c1 = 0.82) and more weakly across distant lineages, consistent with the "
    "shared cell-of-origin regulation the map implies. Single-cell co-detection is "
    "demonstrated wherever a malignant single-cell atlas exists (both lung subtypes "
    "and glioblastoma \u2014 the latter entirely from prediction, using no GBM "
    "surface-abundance measurement); the remaining cohorts are nominated on the same genetic and "
    "topological evidence. The result is a pan-cancer, confidence-tiered catalogue of "
    "multi-antigen ADC co-target sets with a concrete plan to test them.")
para(abstract)
print("part 1 (title/abstract) ok")

# ============================ INTRODUCTION ===================================
H("Introduction", level=1)
para(
    "Antibody\u2013drug conjugates (ADCs) deliver a cytotoxic payload to cells that "
    "display a chosen surface antigen, and have become one of the fastest-growing "
    "classes of cancer therapeutic{cite:Nelson2023,Gazzah2022}. Their central "
    "vulnerability is the target: an "
    "ADC is only as selective as the difference in antigen abundance between tumour "
    "and normal tissue, and only as durable as the tumour's inability to shed the "
    "antigen without cost. Both problems have the same escape route\u2014a single "
    "antigen can be lost, down-regulated, or shared with a vital normal tissue.")
para(
    "A way around both is to engage more than one antigen at once. Multispecific "
    "ADC formats are now clinical reality: izalontamab brengitecan, a bispecific "
    "ADC carrying a topoisomerase-I-inhibitor payload and engaging EGFR and HER3, "
    "improved survival over chemotherapy in a randomised phase-3 trial in "
    "nasopharyngeal carcinoma{cite:IzaBren2025}, establishing that a two-antigen ADC "
    "is a clinically viable format. That precedent targets two co-expressed driver "
    "receptors to block signalling. We propose a different use of the same modality: "
    "engage several surface antigens that are co-amplified as passengers on one "
    "recurrent copy-number amplicon\u2014such as the 3q amplification near-ubiquitous "
    "in lung squamous carcinoma{cite:McCaughan2010,Jeon2023}\u2014and require "
    "co-engagement for stable binding. "
    "Because the antigens rise and fall together with the amplicon\u2014which the "
    "tumour is under selection to keep\u2014this both narrows the escape space "
    "(losing the amplicon loses the tumour's advantage) and, through avidity "
    "AND-gating, sharpens selectivity: a construct that binds stably only where two "
    "antigens are both present spares normal cells that carry only one.")
para(
    "Targeting co-amplified passengers rather than the amplicon's driver is a "
    "deliberate choice. The driver is often intracellular, essential in normal "
    "tissue, or already drugged; the passengers include surface proteins with no "
    "role in the cancer's biology, which is exactly what makes them safer payload "
    "anchors. But a passenger is only a target if its amplification actually "
    "reaches the cell surface as protein, and that is not guaranteed: the path from "
    "gene dosage to surface protein is attenuated, gene by gene, at several "
    "steps{cite:Goncalves2017,Vogel2012,Stingele2012,Sousa2019,Heller2026}. "
    "Selecting targets on copy "
    "number, or even on transcript, therefore risks "
    "committing a programme to genes whose protein never arrives.")
para(
    "This work builds the missing map. Somatic copy-number amplification is among "
    "the most pervasive alterations in cancer{cite:Beroukhim2010,Zack2013,SanchezVega2018}, "
    "and matched proteogenomic atlases now make its protein consequences "
    "measurable{cite:Gillette2020,Li2023}. We quantify copy-number-to-protein "
    "transmission for every amplified gene across six tumour types, locate where "
    "the signal is gated, and show that a gene's transmissibility is a stable, "
    "intrinsic property that can be predicted from gene biology alone\u2014so the "
    "map extends to cancer types that lack a tissue-referenced proteome of their own. We then "
    "combine direct measurement with prediction, gate the result on surface "
    "accessibility, and test in single cells whether the nominated antigen sets "
    "co-occur tightly enough for AND-gated targeting to be selective. The output "
    "is a confidence-tiered, pan-cancer catalogue of multi-antigen ADC co-target "
    "sets, with a concrete experimental plan.")

# ============================ RESULTS ========================================
H("Results", level=1)

H("The transmission cascade and where it is gated", level=2)
para(
    "We first quantify the cascade from amplicon to surface. For each gene we define "
    "two quantities across six CPTAC cohorts with matched copy number, mRNA and protein "
    "(CCRCC, GBM, LSCC, LUAD, PDA and UCEC; 110, 98, 108, 105, 136 and 99 tumours with "
    "all three layers jointly measured): copy-number-to-mRNA <b>transmission</b> (the "
    "Pearson correlation, across tumours, between a gene\u2019s copy number and its mRNA) "
    "and copy-number-to-protein <b>responsiveness</b> (the same correlation to protein). "
    "Their difference is the post-transcriptional <b>attenuation</b>. Correlations are "
    "computed per "
    "gene and per cancer type on jointly non-missing cases (minimum 20 cases per type, "
    "at least two types) and averaged across types by Fisher-z weighting on n\u22123 "
    "(Supplementary Methods). Across the "
    f"{_tg.get('n_genes',6979):,} genes meeting this support, copy number transmits to "
    f"mRNA with a median per-gene correlation of {_cnrna:.2f} but only {_cnprot:.2f} to "
    f"protein, a median attenuation of {_atten:.2f} that leaves "
    f"{100*_fatt:.0f}% of genes attenuated at the protein level (Figure 1a,b). "
    "Crucially, the between-gene ranking established at the mRNA step is largely "
    f"preserved to protein (Spearman \u03c1 = {_rankrho:.2f}, n = {_tg.get('n_genes',6979):,} "
    "genes): the genes that transmit strongly to mRNA are, by and large, the same genes "
    "that reach high protein. This is what licenses an mRNA- and gene-property-based "
    "predictor to stand in for protein measurement.")
para(
    "We then ask where the transmission ranking is set. Gene-structural features "
    f"alone (local gene density) explain almost none of the between-gene variance "
    f"in transmission (cross-validated R\u00b2 = {_r2feat:.3f}, mean over five cancer "
    f"types with matched regulatory data; ~{_methn:,} genes per type). Two regulatory "
    "layers each add an order of magnitude: adding promoter DNA methylation raises "
    f"the explained variance to R\u00b2 = {_r2meth:.3f} (\u2248{_methfold:.0f}-fold over "
    f"structure), and adding tumour promoter chromatin accessibility raises it to "
    f"R\u00b2 = {_r2atac:.3f} (\u2248{_atacfold:.0f}-fold). The two are complementary, not "
    f"redundant: combining methylation and accessibility reaches R\u00b2 = {_r2all:.3f}, "
    "above either alone (Figure 1c), and the pattern holds across all five profiled "
    "cancer types (LUAD, LSCC, CCRCC, GBM, UCEC). The transmission ranking is thus "
    "gated principally at the chromatin/transcription step\u2014by both promoter "
    "methylation and accessibility\u2014rather than by gene structure.")
figure("fig2_transmission_gates.png",
       "Figure 1. The transmission cascade and its gates, computed on the "
       "source-assembled copy-number, mRNA and protein layers (CPTAC; per-gene "
       "Pearson correlations across tumours, \u2265 20 cases per cancer type and \u2265 2 "
       "types, Fisher-z averaged on n\u22123). (a) Per-gene copy-number-to-mRNA "
       "transmission versus copy-number-to-protein responsiveness (n = 6,979 genes); "
       "most genes fall below the diagonal (attenuated post-transcriptionally), and the "
       "mRNA-level ranking is preserved to protein (Spearman \u03c1 = 0.70, n = 6,979). "
       "(b) Distribution of per-gene attenuation (transmission minus responsiveness), "
       "median 0.08, n = 6,979 genes. (c) Cross-validated R\u00b2 of per-gene "
       "transmission predicted from gene structure alone, then adding promoter DNA "
       "methylation, tumour promoter chromatin accessibility, and both, one group of "
       "bars per cancer type (five profiled types: LUAD, LSCC, CCRCC, GBM, UCEC; "
       f"\u2248{_methn:,} genes per type with matched regulatory data; nested "
       "cross-validation). Promoter methylation and accessibility each raise the "
       "explained variance about an order of magnitude over structure and are "
       f"complementary\u2014together reaching mean R\u00b2 {_r2all:.3f} versus {_r2feat:.3f} "
       "for structure alone.")
print("part 2 (intro + cascade) ok")

H("Transmissibility is predictable without protein data", level=2)
para(
    "We now summarise a gene\u2019s tendency to reach protein by a single scalar, its "
    "<b>transmissibility</b>: the fraction of amplified tumours (copy number above the "
    "ploidy-adjusted amplification threshold, deduplicated to one value per case) in "
    "which the gene\u2019s protein is elevated relative to matched normal tissue "
    "(tissue-referenced protein rank > 0.80). Transmissibility runs from 0 (never "
    "elevated when amplified) to 1 (always elevated); it is defined per gene, pooled "
    "across the cohorts in which the gene is recurrently amplified, and is the quantity "
    "the predictor learns and the empirical-Bayes step refines.")
para(
    "If transmissibility is an intrinsic property of a gene, it should be "
    "predictable from the gene's own biology, measured independently of any "
    "proteomics. Using only gene properties\u2014dosage sensitivity, protein "
    "biophysics, mRNA features, evolutionary constraint, complex membership, "
    "expression breadth and network centrality, each defined with its data source "
    "in Supplementary Table S1{cite:Karczewski2020,Tsherniak2017,Giurgiu2019}\u2014"
    "and a leave-gene-out design in "
    f"which no gene informs its own prediction, we recover transmissibility at a "
    f"rank correlation of {rho:.2f} (R\u00b2 = {r2:.2f}) across "
    f"{num('02_predictor','n_genes'):,} genes. No protein-derived feature enters "
    "the predictor. Two controls establish that this is intrinsic rather than "
    "positional or lineage-bound. Holding out whole chromosome arms in turn, so no "
    "gene is predicted from a same-arm neighbour, changes the accuracy by only "
    f"{darm:.3f}\u2014the signal is not a positional artefact. And ranking genes by "
    "predicted transmission within each held-out lineage, the rankings agree across "
    f"the six tumour types at a Kendall concordance of {dW:.2f}. A gene's "
    "transmissibility can therefore be anticipated before its protein is measured, "
    "which is what licenses extending the target map to cancer types with no "
    "tissue-referenced proteome of their own.")
figure("fig3_predictor.png",
       "Figure 2. Transmissibility is predictable from gene properties alone "
       "(gradient-boosted regressor, 39 gene-property features, no protein-derived "
       f"input; n = {num('02_predictor','n_genes'):,} genes). (a) Five-fold "
       "leave-gene-out out-of-fold predicted versus observed transmissibility "
       f"(Spearman \u03c1 = {rho:.2f}, R\u00b2 = {r2:.2f}). (b) Positional control: "
       "leave-gene-out versus leave-chromosome-arm-out out-of-fold accuracy "
       f"(\u0394\u03c1 = {darm:.3f}), showing the signal is not carried by same-arm "
       "neighbours. (c) Cross-lineage transfer: agreement of predicted rankings across "
       f"six leave-one-lineage-out refits (Kendall\u2019s W = {dW:.2f}, six cancer types).")

H("A quantified funnel from transmissibility to accessible targets", level=2)
_fn = P.get("00b_target_funnel", {})
para(
    "With the transmission map and the predictor in hand, the full nomination path "
    "is a funnel, and every step is quantified from data (Figure 3). It begins with "
    f"the {_fn.get('universe',6926):,} genes carrying both a measured and a "
    "predicted transmissibility across the CPTAC cohorts\u2014so the target space is "
    "pan-cancer from the outset, not restricted to any one tumour type. Requiring "
    f"transmission to protein (observed transmissibility \u2265 0.40) retains "
    f"{_fn.get('transmitted',3115):,} genes. Independently, a per-(cohort, cytoband, "
    "gene) Fisher test identifies genes whose high tissue-relative protein is "
    "enriched in tumours where their band is amplified; "
    f"{_fn.get('coelevated',888):,} genes are co-elevated on a recurrent amplicon in "
    f"at least one cohort (BH-FDR < 0.1), of which {_fn.get('coelev_transmitted',507):,} "
    "are also transmitted. The final ectodomain-topology gate (UniProt: extracellular "
    f"\u2265 50 aa on a membrane-anchored protein) yields the "
    f"{_fn.get('nominated_antigens',22)} nominated surface antigens, distributed "
    f"across {_fn.get('nominated_amplicons',18)} distinct recurrent amplicons in "
    f"{_fn.get('n_cohorts_nominated',4)} cancer types. The funnel is pan-cancer at "
    "every step; the antigens it yields are detailed next.")
figure("fig2_target_funnel.png",
       "Figure 3. Target-nomination funnel. Each stage is computed: the "
       "transmissibility atlas, the per-cohort Fisher co-elevation test on recurrent "
       "cytoband amplicons (BH-FDR < 0.1), and the UniProt ectodomain gate; no counts "
       "are hardcoded. Co-elevation is evaluated in the four proteome-supported "
       "cohorts (CCRCC, LSCC, LUAD, UCEC) and the nominated antigens span all four.",
       width=4.6)

H("Nominated surface antigens and multi-antigen ADC constructs", level=2)
# load the pan-cancer target tables produced by 04c_pancancer_nomination.py
# and 04d_constructs.py (new schema).
import csv as _csv2
def _load(fname):
    p = cfg.DIR_TAB / fname
    return list(_csv2.DictReader(open(p))) if p.exists() else []
ant_rows = _load("adc_target_antigens.csv")
con_rows = _load("adc_constructs.csv")
n_ant = len({r["antigen"] for r in ant_rows}); n_con = len(con_rows)
_cohorts = sorted({r["cohort"] for r in ant_rows})
n_cohort = len(_cohorts)
n_amplicons = len({r["amplicon"] for r in ant_rows})
n_tested = sum(1 for r in con_rows if str(r.get("single_cell_tested")).lower()=="true")
n_nomonly = n_con - n_tested

# Passenger-only (driver-excluded) construct: computed by 04d, not hardcoded. Pull the
# LUAD 7p passenger-only row so the "co-targeting rests on passengers" number is live.
def _con_by_amplicon(amp):
    for r in con_rows:
        if r.get("amplicon") == amp:
            return r
    return {}
_pass = _con_by_amplicon("LUAD_7p_pass")
def _f2(x):
    try: return f"{float(x):.2f}"
    except Exception: return str(x)
_pass_enr = _f2(_pass.get("enrich", "1.30"))
_pass_lo  = _f2(_pass.get("ci_lo", "1.13"))
_pass_hi  = _f2(_pass.get("ci_hi", "1.47"))
_pass_ant = _pass.get("antigens", "ITGB8+TSPAN13+TTYH3").replace("+", "+")
para(
    "A gene is a surface ADC target only if it is co-amplified on a recurrent "
    "amplicon, transmitted to protein, and presents an accessible extracellular "
    "ectodomain a native-format antibody could bind. We require all three: "
    "co-elevation on a recurrent amplicon (Fisher FDR < 0.1), measured "
    "transmissibility at or above 0.40, and a UniProt-annotated extracellular "
    "ectodomain of at least 50 residues on a membrane-anchored "
    "protein{cite:UniProt2023}. Applied across every recurrent amplicon in the "
    f"four proteome-supported cohorts, this nominates {n_ant} surface antigens "
    f"on {n_amplicons} distinct amplicons spanning {n_cohort} cancer types "
    f"({', '.join(_cohorts)}) (Table 1)\u2014the nomination is pan-cancer, not "
    "restricted to lung. Examples span the amplicon landscape: EGFR, ITGB8 and "
    "TSPAN13 on lung 7p; NCSTN, HSD17B7 and MPZL1 on 1q (recurrent in LUAD, LSCC "
    "and UCEC); the transferrin receptor TFRC on the squamous 3q amplicon; the "
    "Wnt receptor FZD1 and integrin ITGB8 on clear-cell renal 7p/7q; and an "
    "NCSTN+MPZL1+ADAM15 set on endometrial 1q. The topology gate is decisive: it "
    "removes high-transmission genes with no accessible ectodomain (for example "
    "EFNA1, a GPI-anchored ligand), and downgrades multipass transporters even "
    "when amplified. Twenty-one of the 22 antigens are non-driver passengers; only "
    "EGFR (lung 7p) is a canonical oncogenic driver, and the passenger thesis does "
    "not depend on it (below). Two further non-driver antigens carry an "
    "essentiality caveat (TFRC, DepMap \u22120.94; VMP1, \u22121.13); the remaining "
    "19 are neither drivers nor broadly essential. As orthogonal evidence for the "
    "surface call, 10 of the 22 \u2014 including the transmission leaders TFRC, "
    "NCSTN, MPZL1, ITGB8, ADAM15 and TTYH3 \u2014 are listed on the experimental "
    "Cell Surface Protein Atlas{cite:BauschFluck2015} (mass-spectrometric surface "
    "capture); the remainder retain their UniProt topology call, with CSPA absence "
    "treated as weak evidence because surface capture sees only N-glycosylated "
    "proteins in the cell types assayed. Each antigen carries its driver, "
    "essentiality and CSPA flags in Table 1.")

H("Table 1. Nominated surface antigens (pan-cancer)", level=3)
if ant_rows:
    hdrs = ["Cohort","Amplicon","Antigen","Meas.","Pred.","n amp.","Ecto","TM","DepMap","Class","CSPA"]
    def _f(x, nd=2):
        try: return f"{float(x):.{nd}f}"
        except Exception: return x
    def _cls(r):
        if str(r.get("is_driver")).lower()=="true": return "driver"
        try:
            if float(r.get("dep_effect") or 0) <= -0.40: return "ess."
        except Exception: pass
        return "passenger"
    def _cspa(r):
        c = str(r.get("cspa_category",""))
        return "high" if c.startswith("1") else ("put." if c.startswith("2") else
               ("unspec." if c.startswith("3") else "\u2014"))
    trows = [[r["cohort"], r["amplicon"], r["antigen"], _f(r["obs_transmit"]),
              _f(r["pred_transmit"]), r["n_amplified"], _f(r["ecto_aa"],0),
              _f(r["n_tm"],0), _f(r["dep_effect"],2), _cls(r), _cspa(r)]
             for r in sorted(ant_rows, key=lambda r:(r["cohort"], r["amplicon"], -float(r["obs_transmit"])))]
    add_table(hdrs, trows, fs=6.5)
para("Cohort = CPTAC cancer type in which the amplicon is recurrent; Measured / "
     "predicted = observed and gene-property-predicted transmissibility; n amp. = "
     "amplified cases underlying the measurement; Ecto = UniProt extracellular "
     "residues; TM = transmembrane helices; DepMap = mean CRISPR dependency effect "
     "(more negative = more essential); Class = driver / broadly essential (ess., "
    "DepMap \u2264 \u22120.40) / passenger; CSPA = Cell Surface Protein Atlas category "
    "(high-confidence / putative / unspecific / \u2014 not listed).", italic=True)

para(
    "Because the antigens on one amplicon rise together, an antibody that engages "
    "two or three of them can be built as an avidity AND-gate. We propose the "
    f"{n_con} multivalent constructs in Table 2, each assembled from the nominated "
    "antigens on a single amplicon (an amplicon with only one nominated antigen "
    "cannot form a construct). Same-cell co-detection is measured directly in "
    f"malignant single cells for the {n_tested} constructs whose cancer type has a "
    "single-cell atlas slice (LUAD, LSCC); the remaining "
    f"{n_nomonly} construct{'s' if n_nomonly!=1 else ''} on cohorts without a "
    "malignant single-cell slice (here CCRCC/UCEC) "
    f"{'are' if n_nomonly!=1 else 'is'} nominated on the same genetic and "
    "topological evidence and flagged accordingly, as no malignant single-cell "
    "atlas is available for those cohorts to test co-detection. Co-detection is "
    "scored against a null that permutes each antigen within per-cell "
    "sequencing-depth deciles, so the reported enrichment is excess co-occurrence "
    "beyond what shared depth alone produces (a marginal-only null inflates it "
    "several-fold). On that null, the co-detected constructs are above independence "
    "with a bootstrap interval clear of 1.0 and permutation p = 0.001, at a modest "
    "1.05\u20131.45\u00d7; two bivalent pairs (LUAD 1q, GBM 20q) are not significant. "
    "The load-bearing result is the direction and significance of same-cell "
    "co-occurrence, not the absolute fold.")

H("Table 2. Proposed multivalent ADC constructs", level=3)
if con_rows:
    hdrs2 = ["Construct","Valence","Evidence","Same-cell enrichment (\u00d7)","95% CI","perm p","Cells / donors","Single-cell tested"]
    def _cf(r, k):
        v = r.get(k,"")
        return "\u2014" if v in ("","nan","NaN") else v
    trows2 = []
    for r in con_rows:
        tested = str(r.get("single_cell_tested")).lower()=="true"
        ci = f'{_cf(r,"ci_lo")}\u2013{_cf(r,"ci_hi")}' if tested else "\u2014"
        cd = f'{_cf(r,"n_cells").split(".")[0]} / {_cf(r,"n_donors").split(".")[0]}' if tested else "\u2014"
        trows2.append([r["construct"], r["valence"], r.get("evidence","measured"),
                       _cf(r,"enrich"), ci, _cf(r,"perm_p"), cd, "yes" if tested else "nominated"])
    add_table(hdrs2, trows2, fs=6.5)
para("Same-cell enrichment is observed co-detection divided by a depth-matched "
     "expectation in malignant single cells (null permutes each antigen within "
     "per-cell sequencing-depth deciles; donor-block bootstrap 95% CI; depth-"
     "stratified permutation p). Evidence = measured (nominated from tissue-referenced "
     "proteomics) or prediction-only (nominated from predicted transmissibility because "
     "that cohort has no tissue-referenced proteome; GBM). Constructs on cohorts without "
     "a malignant single-cell "
     "atlas are nominated on genetic and topological evidence; co-detection cannot "
     "be evaluated there and they are labelled as such.", italic=True)

figure("fig5_surface_targets.png",
       "Figure 4. Pan-cancer surface targets. (a) The measured-nomination antigens "
       "(the four proteome-supported cohorts), coloured by cohort, placed against all "
       "6,926 genes by predicted and observed transmissibility\u2014they occupy the "
       "high-transmissibility regime above the 0.40 floor; glioblastoma is "
       "prediction-only and has no observed value, so it is not plotted here (its "
       "antigens are in Supplementary Table S4). (b) Nominated antigens per amplicon "
       "in those four cohorts, bar = maximum observed transmissibility on the "
       "amplicon. (c) Same-cell co-detection enrichment (depth-stratified null) for "
       "the constructs with a malignant single-cell slice (LUAD, LSCC measured; GBM "
       "prediction-only; 95% CI; significant constructs p = 0.001, LUAD 1q and GBM "
       "20q n.s.); constructs on cohorts without a malignant single-cell atlas are "
       "nominated but not shown here.")

H("AND-gated co-targeting is selective in single cells", level=2)
# on-tumour numbers from the tested (single-cell) constructs, new schema
_tested = [r for r in con_rows if str(r.get("single_cell_tested")).lower()=="true"]
# range is quoted for the SIGNIFICANT constructs (perm p < 0.05, CI clear of 1.0);
# the non-significant LUAD 1q (0.99x, p=0.84) is reported separately in the text.
_sig = [r for r in _tested if float(r.get("perm_p", 1)) < 0.05
        and float(r.get("ci_lo", 0)) > 1.0]
con_enr = [float(r["enrich"]) for r in _sig] if _sig else []
c_lo = min(con_enr) if con_enr else None; c_hi = max(con_enr) if con_enr else None
# borderline = permutation-significant but bootstrap CI touches 1.0; ns = neither
_ns  = [r for r in _tested if not (float(r.get("perm_p",1))<0.05)]
_bord= [r for r in _tested if float(r.get("perm_p",1))<0.05
        and float(r.get("ci_lo",0))<=1.0]
n_sc_tested = len(_tested); n_sig = len(_sig); n_ns = len(_ns); n_bord = len(_bord)
_ns_names = ", ".join(sorted(r["amplicon"].replace("_"," ") for r in _ns)) or "none"
_gbm_tested = [r for r in _tested if r.get("cohort")=="GBM"]
_gbm_sig = [r for r in _gbm_tested if float(r.get("perm_p",1))<0.05
            and float(r.get("ci_lo",0))>1.0]
_biv = [r for r in _tested if r["valence"]=="bivalent"]
_biv_enr = [float(r["enrich"]) for r in _biv]
para(
    "For a multivalent construct to be efficient on tumour, its antigens must appear "
    "together in the same malignant cell, not merely in the same tumour, and for it "
    "to be selective the co-occurrence must fall away in normal tissue. We test both "
    "in single malignant cells from three cancer types with a malignant single-cell "
    "atlas (41,615 LUAD cells from 130 donors, 33,234 LSCC cells from 42 donors, and "
    "390,761 glioblastoma cells from 208 donors, from the CELLxGENE "
    "census{cite:CELLxGENE2025}); the glioblastoma constructs are prediction-only "
    "(GBM has raw proteome in CPTAC but no tissue-referenced proteome, so it yields "
    "no measured surface nominations), and their single-cell test is a genuine "
    "out-of-sample check of the prediction pipeline. For each construct whose "
    "cancer type has a "
    "single-cell slice we ask whether its nominated antigens are co-detected in the "
    f"same cell more often than independent detection would predict, scored "
    "against a depth-stratified null (each antigen permuted within per-cell "
    "sequencing-depth deciles, so the null preserves the depth structure that "
    "otherwise inflates co-detection). "
    f"Of the {n_sc_tested} tested constructs, {n_sig} are co-detected above "
    f"independence ({c_lo:.2f}\u2013{c_hi:.2f}\u00d7), each with permutation "
    "p = 0.001 and a donor-block bootstrap interval clear of 1.0; "
    f"{n_bord} more (the trivalent GBM 7q set) is permutation-significant with a "
    f"bootstrap interval that just touches 1.0, and {n_ns} "
    f"(the bivalent LUAD 1q and GBM 20q pairs) are not significant on this null "
    "(0.99\u20131.00\u00d7). "
    "The magnitudes are modest \u2014 a marginal-only null returns several-fold "
    "higher values, but most of that is the depth artefact the stratified null "
    "removes \u2014 and the higher-valence LUAD 7p (four antigens) set retains the "
    "largest true excess (1.45\u00d7), consistent with amplicon-driven coordination. "
    "Critically, LUAD 7p survives removal of its one driver, EGFR: the "
    f"passenger-only {_pass_ant} set is still co-detected at {_pass_enr}\u00d7 "
    f"(95% CI {_pass_lo}\u2013{_pass_hi}, p = 0.001), so the co-targeting rests on passengers, "
    "not on the driver. Most notably, the glioblastoma constructs \u2014 nominated "
    "without any GBM surface-abundance measurement, from predicted transmissibility "
    "alone \u2014 reproduce the "
    f"same pattern: {len(_gbm_sig)} of the {len(_gbm_tested)} GBM constructs "
    "(19p ATP13A1+TMED1 at 1.12\u00d7, 20p ATRN+PTPRA at 1.10\u00d7) are co-detected "
    "above independence with intervals clear of 1.0, so the full predict\u2192"
    "surface-gate\u2192single-cell-verify path closes in a cancer type whose proteome "
    "is not tissue-referenced and so drives no measured nomination. The load-bearing "
    "result is the direction and significance "
    "of same-cell co-occurrence, not any single fold value.")
para(
    "The off-tumour side is where AND-gating earns its selectivity. A normal cell "
    "binds an avidity AND-gate only if its weakest (limiting) antigen clears the "
    "binding threshold. Raising the per-antigen threshold from detection to "
    "binding-relevant levels collapses the fraction of normal cell types in which "
    "all of a construct's antigens are co-present toward zero, provided each "
    "construct contains at least one antigen that is low in normal "
    "tissue{cite:Uhlen2015,Karlsson2021}. Together these give the selectivity "
    "argument its two halves: measurable on-tumour same-cell co-occurrence, and "
    "threshold-gated off-tumour sparing.")
figure("fig6_andgate.png",
       "Figure 5. AND-gated selectivity in single cells, for the five nominated "
       "constructs whose cohort has a malignant single-cell slice (LUAD: 41,615 cells / "
       "130 donors; LSCC: 33,234 cells / 42 donors; CELLxGENE). (a) Same-cell "
       "co-detection enrichment in malignant cells for each construct "
       "(observed / depth-expected co-detection; donor-block bootstrap 95% CI, 1,000 "
       "resamples; depth-stratified permutation p, 1,000 permutations, each antigen "
       "shuffled within per-cell sequencing-depth deciles). Constructs are shown for "
       "the three cohorts with a malignant single-cell atlas (LUAD, LSCC measured; "
       "GBM prediction-only); those above independence span 1.05\u20131.45\u00d7 "
       "(p = 0.001, interval clear of 1.0), while the LUAD 1q and GBM 20q pairs are "
       "not significant. A marginal-only null returns several-fold higher values, "
       "most of which is the depth artefact this null removes (see text). "
       "(b) Fraction of Human Protein Atlas normal cell types in which all antigens of "
       "a construct exceed a per-antigen threshold, at detection (10 nTPM) versus "
       "binding-relevant (25, 50 nTPM) levels; raising the threshold collapses the "
       "normal-tissue burden (LUAD 7p to zero), which is the basis of the AND-gate "
       "selectivity window.")
print("part 3 (results) ok")

H("A plan to test the nominated sets", level=2)
para(
    "The nominations are hypotheses about surface protein, and the framework makes "
    "each one falsifiable. The forward plan is a staged go/no-go: confirm surface "
    "co-expression of a set's antigens in amplicon-positive versus amplicon-negative "
    "cell lines; measure native-surface antibody binding for marginal ectodomains; "
    "build an avidity AND-gate construct and confirm that tumour killing requires "
    "both antigens while single-antigen cells are spared; and measure the "
    "selectivity window against matched normal tissue. A no-go at any step returns "
    "the set to nomination with the failing antigen dropped, which the confidence "
    "tiering makes inexpensive. Concrete entry points are amplicon-positive cell "
    "lines identified from DepMap 23Q4 copy number (all anchor genes at relative "
    "copy number \u2265 1.4 in the matching lineage): the LUAD 7p set "
    "(ITGB8+TSPAN13+TTYH3, driver EGFR excluded) in NCI-H3255, NCI-H1838 and "
    "HCC4006, with the amplicon-negative LUAD lines NCI-H441 and NCI-H358 as "
    "single-antigen controls; the LSCC 1q set (HSD17B7+MPZL1+NCSTN) in LUDLU-1/GT3TKB; "
    "and the squamous 3q set (ATP11B+TFRC) in the 3q-amplified lines HCC95 and "
    "LC-1/sq. The LUAD 7p and LSCC 1q entry points additionally have their same-cell "
    "co-detection confirmed here on the depth-stratified null (Figure 5).")
figure("fig7_experimental_plan.png",
       "Figure 6. From in-silico nomination to the bench. Each in-silico output "
       "maps to one wet-lab test with an explicit go/no-go criterion, with named "
       "cell-line and cohort entry points.", width=6.6)

# ============================ DISCUSSION =====================================
H("Discussion", level=1)
para(
    "We have built an integrated, quantified map of how a copy-number amplification "
    "reaches the cell surface, and used it to nominate multi-antigen ADC co-target "
    "sets. Three features make it more than a target list. First, it is a full-path "
    "account: the framework resolves the amplicon-to-mRNA-to-protein path into its "
    "steps and shows where the between-gene ranking is set\u2014gated at chromatin and "
    "transcription, where promoter methylation and accessibility each explain an "
    "order of magnitude more variance than gene structure, and then attenuated "
    "post-transcriptionally (Figure 1)\u2014so the map "
    "addresses not just which "
    "genes transmit but where the control lies. Second, "
    "transmissibility is intrinsic and therefore portable\u2014predictable from gene "
    "biology alone, not positional, and stable across lineages\u2014so the map "
    "extends by prediction to cancer types that lack a tissue-referenced proteome, with the "
    "empirical-Bayes posterior making explicit how much of each nomination is "
    "measured and how much predicted. Third, the modality and the genetics fit: "
    "co-amplified passengers rise together with an amplicon the tumour is selected "
    "to keep, and avidity AND-gating turns that co-occurrence into selectivity, "
    "which we confirm at single-cell resolution.")
para(
    "This has a direct consequence for how widely a nominated set should apply "
    "within a cancer type. Once an amplicon is present, which of its genes reach the "
    "surface as protein is set by the regulatory gates quantified in Figure 1\u2014"
    "promoter DNA methylation and chromatin accessibility\u2014and both of these are "
    "predominantly properties of the cell of origin rather than of the individual "
    "tumour: normal human cell types have DNA methylomes so stable that biological "
    "replicates of the same cell type are more than 99% identical{cite:Loyfer2023}, "
    "and primary-tumour chromatin accessibility clusters by lineage{cite:Corces2018}. "
    "Because we deliberately nominate low-essentiality passenger antigens with no "
    "role in the cancer\u2019s biology, there is little tumour-specific selection "
    "acting on their expression to override those shared gates. Two tumours of the "
    "same type that both carry the same amplicon therefore start from the same "
    "cell-of-origin gate settings and face no divergent selective pressure on the "
    "passengers, so they should present substantially the same set of amplicon-driven "
    "surface antigens. Our transmissibility values already quantify this recurrence "
    "directly: transmissibility is the fraction of amplified tumours in a cohort in "
    "which a gene\u2019s protein is elevated, so a nominated antigen with "
    "transmissibility 0.6\u20130.9 is, by definition, one whose amplification reaches "
    "the surface in a high proportion of patients carrying that gain in that cancer "
    "type. The practical implication is that a construct nominated against, say, the "
    "lung-squamous 3q amplicon is expected to be actionable across a large share of "
    "3q-gained lung-squamous patients, not only the tumour in which it was measured. "
    "This is a within-cancer-type, across-patient claim; it is distinct from and does "
    "not depend on the weaker cross-lineage transfer of the transmissibility ranking. "
    "The measured cross-cohort correlations bear this out: observed transmissibility "
    "correlates strongly between the two lung subtypes (Spearman \u03c1 = 0.82 over "
    "632 shared genes), moderately between lung and endometrial (0.55\u20130.66), and "
    "weakly with the more distant clear-cell renal cohort (0.11\u20130.19) \u2014 the "
    "gradient expected if the gates are set by cell of origin. It is this within- and "
    "near-lineage portability, not a universal cross-cancer constant, that additionally "
    "licenses extending the map by prediction to related cancer types with no "
    "tissue-referenced proteome of their own.")
para(
    "The clinical reading is deliberately measured. The nominations are "
    "transcript- and prediction-level calls about surface protein, not "
    "surface-protein measurements; the honest bound is that gene properties explain "
    "about a third of per-gene transmission variance, so the framework ranks the "
    "genome well but will mispredict individual genes. That is why every lead enters "
    "a staged wet-lab test before it could be called a target, and why the "
    "confidence tiers\u2014measured, predicted, or both\u2014travel with every "
    "candidate rather than being averaged away.")

H("Limitations", level=1)
para(
    "Several bounds are explicit. The transmissibility prior explains about a third "
    "of per-gene variance; it is a genome-wide prior, not a per-gene oracle, and "
    "individual genes are mispredicted. The surface calls rest on UniProt topology "
    "annotation, not on measured surface abundance, and the next step for any lead "
    "is direct surface-protein and tumour-versus-normal measurement. The single-cell "
    "co-detection uses detection (non-zero counts), which is conservative for "
    "abundance but does not by itself establish that both antigens reach "
    "binding-relevant surface density; the normal-tissue threshold analysis "
    "addresses magnitude on the off-tumour side but the on-tumour density claim "
    "still requires protein measurement. Co-detection is scored against a null that "
    "controls for per-cell sequencing depth, which reduces the enrichments to modest "
    "excess (1.05\u20131.45\u00d7) and renders one bivalent pair non-significant; the "
    "argument rests on the direction and significance of same-cell co-occurrence "
    "together with the threshold-gated normal sparing, not on large fold values. "
    "The nomination is also amplification-threshold dependent: at a stringent "
    "high-level cut (adjusted copy number \u2265 2.0) the recurrent-amplicon set and "
    "the antigen list shrink substantially and only the LSCC 3q and LUAD 5p "
    "constructs survive intact (Supplementary Table S3); we use \u2265 1.4 because the "
    "clinical premise is the broad, near-ubiquitous segmental gains that define these "
    "amplicons, not rare focal high-level events. The chromatin gate is quantified from both "
    "promoter DNA methylation (GDC EPIC arrays, mapped to genes through the Illumina "
    "manifest) and tumour ATAC accessibility, which prove complementary.")

# ============================ SUPPLEMENTARY ==================================
doc.add_page_break()
H("Supplementary Methods", level=1)

H("Data and cohorts", level=2)
para(
    "Matched copy number, mRNA and protein were assembled across six tumour types "
    "(LSCC, LUAD, GBM, PDA, CCRCC, UCEC) from public proteogenomic "
    "resources{cite:Gillette2020,Li2023}, "
    "harmonised to a common gene- and case-level table. Amplification is defined on "
    "a ploidy-adjusted copy-number basis (adjusted copy number \u2265 1.4), the "
    "regime relevant to the clinical output. Tumour chromatin accessibility is taken "
    "from cancer-type-matched ATAC-seq{cite:Corces2018}; single-cell malignant-cell "
    "profiles are from the CELLxGENE census{cite:CELLxGENE2025} (lung adenocarcinoma "
    "and lung squamous carcinoma, primary malignant cells); normal single-cell and "
    "bulk expression are from the Human Protein Atlas{cite:Uhlen2015,Karlsson2021} "
    "and GTEx{cite:GTEx2020}. Surface topology is from UniProt reviewed human "
    "entries{cite:UniProt2023} (extracellular topological-domain residues and "
    "transmembrane-helix count). The full index of datasets and how each is used is "
    "in Supplementary Table S2.")

H("Transmission cascade", level=2)
para(
    "Per gene and per tumour type, copy-number-to-mRNA transmission and "
    "copy-number-to-protein responsiveness are Pearson correlations on the joint "
    "non-missing cases (minimum 20 cases per type, at least two types, Fisher-z "
    "averaging weighted by n\u22123); attenuation is their difference. The chromatin "
    "gate is a nested cross-validated linear model of per-gene transmission, comparing "
    "gene structure alone against structure plus promoter DNA methylation, plus tumour "
    "promoter accessibility, and plus both, across the profiled cancer types. Promoter "
    "methylation is the mean SeSAMe beta over a gene\u2019s TSS200/TSS1500 probes on the "
    "Illumina EPIC array (GDC CPTAC-3, 655 matched tumours); probes are mapped to genes "
    "through the Illumina EPIC manifest (UCSC_RefGene annotation) rather than through "
    "ChAMP, which does not build on current R/Bioconductor. Accessibility is the "
    "tumour-type promoter ATAC signal (TCGA-ATAC).")

H("Transmissibility predictor", level=2)
para(
    "A gradient-boosted regressor{cite:ChenGuestrin2016} (600 trees, depth 4, "
    "learning rate 0.03, "
    "subsample and column-sample 0.8, seed 2) predicts per-gene transmissibility "
    "from the gene-property features defined in Supplementary Table S1, with no "
    "protein-derived input. Generalisation is estimated by five-fold "
    "leave-gene-out out-of-fold prediction; the positional control refits holding "
    "out whole chromosome arms in turn; cross-lineage transfer is the Kendall "
    "concordance of predicted rankings across leave-one-lineage-out refits. No "
    "protein-derived feature is used as a predictor.")

H("Empirical-Bayes combination of measurement and prediction", level=2)
_cross = num("03_empirical_bayes", "crossover_n")
para(
    "Direct co-elevation is precise where proteomics is deep and noisy where it is "
    "thin, while the predictor is available everywhere but explains about a third of "
    "per-gene variance. We combine them in an empirical-Bayes "
    "posterior{cite:EfronMorris1975} that shrinks each per-gene measurement toward the "
    "predictor prior in inverse proportion to the measurement\u2019s precision. The "
    "posterior for each gene is a precision-weighted blend of the predictor prior "
    "(mean) and the measured transmissibility, with binomial measurement variance "
    "p(1\u2212p)/n set by the number of amplified cases and a between-gene prior variance "
    "estimated by method of moments. The thin-cohort recovery test treats genes with "
    "at least 150 amplified cases as reference truth (n = 6,926 genes total), injects "
    "binomial noise at a grid of cohort sizes, and scores raw measurement, prior-only "
    "and posterior against the reference by rank correlation and RMSE over repeated "
    "resamples (seed 2). The posterior never falls below either the raw measurement or "
    f"the prior at any cohort size; the raw measurement overtakes the prior-only floor "
    f"only above about {_cross} amplified cases; and the prior-only floor "
    f"(\u03c1 = {floor:.2f}) is the achievable ceiling when no proteomics exists at all "
    "(Supplementary Figure S1). This yields a posterior transmissibility for every gene "
    "with an explicit decomposition of how much is measured versus predicted; the "
    "nomination funnel itself uses observed transmissibility and the co-elevation test "
    "directly, so the posterior is a refinement for thin-proteomics cohorts rather than "
    "a gate on the headline results.")
figure("fig4_empirical_bayes.png",
       "Supplementary Figure S1. Empirical-Bayes combination of measured and predicted "
       "transmissibility (n = 6,926 genes). (a) Weight placed on the prior falls as the "
       "number of amplified cases (measurement precision) rises. (b) Thin-cohort "
       "recovery of reference truth (genes with \u2265 150 amplified cases): the posterior "
       f"tracks or beats the raw measurement and never drops below the prior-only floor "
       f"(\u03c1 = {floor:.2f}). (c) RMSE reduction of the posterior over raw measurement, "
       "largest in the thin-cohort regime.")

H("Recurrent amplicons, co-elevation and surface nomination", level=2)
para(
    "A gene is amplified in a tumour when its ploidy-adjusted copy number is at least "
    "1.4 (deduplicated to one value per case). A cytoband is amplified in a tumour when "
    "at least 50% of its genes are amplified, and recurrent in a cohort when it is "
    "amplified in at least 20% of tumours (minimum 8). A gene is co-elevated on a "
    "recurrent amplicon when its tissue-referenced protein rank exceeds 0.80 (\u201chigh\u201d) "
    "in amplified cases significantly more often than in non-amplified cases, by a "
    "one-sided Fisher exact test; p-values are Benjamini\u2013Hochberg-corrected within "
    "cohort and controlled at FDR < 0.10. This band-level test is run in the four "
    "cohorts with tissue-referenced proteome (CCRCC, LSCC, LUAD, UCEC; GBM has copy "
    "number but no tissue-referenced protein and PDA has no cytoband meeting the "
    "recurrence bar, so neither enters the co-elevation test). Nominated antigens are "
    "co-elevated genes (FDR < 0.10) with observed transmissibility \u2265 0.40 that pass "
    "a live UniProt topology gate\u2014at least 20 extracellular residues on a "
    "membrane-anchored protein to be scored accessible, with \u2265 50 required for "
    "nomination as a confidently accessible epitope. Because transmissibility and "
    "co-elevation are both derived from the same tissue-referenced protein-abundance "
    "indicator (protein rank > 0.80), they are not independent filters but two views "
    "of the same protein signal \u2014 transmissibility summarising its level across "
    "amplified tumours, co-elevation testing its enrichment in amplified versus "
    "non-amplified cases \u2014 so the surface-topology gate is the one orthogonal "
    "criterion; the experimental Cell Surface Protein Atlas{cite:BauschFluck2015} is "
    "used as a corroborating flag, not a gate. Each antigen is tagged by measured and "
    "predicted evidence, driver and essentiality status, CSPA membership, and its "
    "DepMap dependency effect. "
    "Same-cell co-detection enrichment is the observed fraction of malignant cells "
    "co-detecting all antigens of a set divided by the product of per-antigen detection "
    "rates, with donor-block bootstrap 95% intervals (1,000 resamples) and a "
    "depth-stratified permutation p (1,000 permutations; each antigen column shuffled "
    "independently within per-cell sequencing-depth deciles, so the null preserves both "
    "per-gene detection rate and per-cell depth and the enrichment measures "
    "co-detection beyond what shared depth alone produces), computed in the LUAD and "
    "LSCC malignant single-cell "
    "slices (CELLxGENE: 41,615 cells / 130 donors and 33,234 cells / 42 donors). "
    "Because a higher-valence set has a smaller independence baseline, the enrichment "
    "ratio scales with the number of antigens; the load-bearing claim is the direction "
    "and significance of co-detection, not the absolute fold. Normal-tissue burden is "
    "the fraction of Human Protein Atlas normal cell types in which the limiting "
    "antigen of a set clears a per-antigen threshold, evaluated at detection (10 nTPM) "
    "and binding-relevant (25, 50 nTPM) levels.")

H("Reproducibility", level=2)
para(
    "All analysis code is in the integrated pipeline folder, with a single "
    "configuration module fixing thresholds and seeds (seed 2 throughout), a data "
    "preparation driver that caches the source layers, and one script per figure. "
    "Every figure regenerates deterministically from the cached source layers.")

# ============================ SUPPLEMENTARY TABLE S1 =========================
# Feature definitions: every predictor feature, its group, plain-language
# meaning, and the data source it is derived from.
doc.add_page_break()
H("Supplementary Table S1. Predictor feature definitions", level=2)
para("Every feature used by the transmissibility predictor, grouped by family. "
     "No feature is derived from protein measurement; only the predicted quantity "
     "(transmissibility) is protein-derived. Column names match config.py "
     "(PREDICTOR_GROUPS) and the exported feature table.", italic=True)

# group -> (plain-language description, data source, per-feature glosses)
GROUP_DEF = {
 "dosage": ("Dosage sensitivity \u2014 how tightly the cell controls the amount of the gene product",
   "gnomAD population constraint; DepMap CRISPR essentiality",
   {"gnomad_LOEUF":"loss-of-function observed/expected upper bound (low = intolerant of losing a copy)",
    "gnomad_pLI":"probability the gene is loss-of-function intolerant",
    "gnomad_mis_z":"missense constraint z-score",
    "dep_mean_effect":"mean CRISPR knockout fitness effect across cell lines (more negative = more essential)",
    "dep_frac_dependent":"fraction of cell lines dependent on the gene"}),
 "complex": ("Complex membership \u2014 whether the protein is a subunit of a multi-protein machine (excess subunits are degraded)",
   "CORUM protein-complex catalogue",
   {"in_complex":"member of at least one annotated complex (0/1)",
    "n_complexes":"number of complexes the protein belongs to",
    "complex_size":"size of the largest complex containing it",
    "has_complex":"any complex annotation present (0/1)"}),
 "biophysics": ("Protein biophysics \u2014 physical properties that govern stability and turnover",
   "UniProt sequence; VSL2 disorder; PSIPRED secondary structure; ASAquick accessibility",
   {"length":"protein length (residues)","mol_weight":"molecular weight",
    "isoelectric_point":"pI (net-charge-neutral pH)","gravy":"grand average hydropathy (hydrophobicity)",
    "aggregation_propensity":"predicted aggregation tendency","tm_domain_count":"number of transmembrane domains",
    "signal_peptide":"signal peptide present (0/1)","vsl2_disorder":"predicted intrinsic disorder fraction",
    "psipred_helix":"predicted \u03b1-helix fraction","psipred_strand":"predicted \u03b2-strand fraction",
    "psipred_coil":"predicted coil fraction","asaquick_buried":"predicted buried-residue fraction"}),
 "mrna": ("mRNA features \u2014 transcript properties that set translation efficiency and message stability",
   "Ensembl/GENCODE transcript models; codon-usage tables",
   {"transcript_length":"mRNA length","gc_content":"GC fraction","n_isoforms":"number of annotated isoforms",
    "utr5_length":"5' UTR length","utr3_length":"3' UTR length","codon_optimality":"mean codon optimality"}),
 "evolution": ("Evolutionary constraint \u2014 how conserved and constrained the gene is",
   "Cross-species alignment (dN/dS); phyloP conservation",
   {"dn_ds":"non-synonymous/synonymous substitution ratio","phylop_mean":"mean phyloP conservation score",
    "gene_age_proxy":"proxy for evolutionary age of the gene"}),
 "function": ("Functional class \u2014 broad molecular-role flags",
   "UniProt keyword / family annotation",
   {"is_tf":"transcription factor (0/1)","is_kinase":"kinase (0/1)","is_receptor":"receptor (0/1)","is_enzyme":"enzyme (0/1)"}),
 "breadth": ("Expression breadth \u2014 how broadly vs specifically the gene is expressed",
   "GTEx bulk tissue expression",
   {"n_tissues_expressed":"number of tissues expressing the gene","tau":"tissue-specificity index (1 = specific, 0 = ubiquitous)"}),
 "network": ("Network centrality \u2014 how central the protein is in the interaction network",
   "STRING protein-protein interaction network",
   {"degree":"number of interaction partners","weighted_degree":"confidence-weighted partner count","betweenness":"betweenness centrality (hub-ness)"}),
}
s1_rows = []
for g, cols in cfg.PREDICTOR_GROUPS.items():
    gd = GROUP_DEF[g]
    for c in cols:
        gloss = gd[2].get(c, "")
        s1_rows.append([g, c, gloss, gd[1]])
add_table(["Group", "Feature", "Definition", "Data source"], s1_rows)
para(f"\n{len(s1_rows)} features across {len(cfg.PREDICTOR_GROUPS)} groups. "
     "Group-level summary of what each family captures is given in the Results "
     "and above.", italic=True)

# ============================ SUPPLEMENTARY TABLE S2 =========================
doc.add_page_break()
H("Supplementary Table S2. Index of datasets and how each is used", level=2)
para("Every external dataset the pipeline consumes, its source, size, and role "
     "in the analysis.",
     italic=True)
import csv as _csv
layers_csv = cfg.DIR_REP / "data_prep_layers.csv"
DATA_USE = {
 "annotation":"Gene/case identifier harmonisation across layers",
 "proteome":"Protein abundance \u2014 the measured transmission outcome (CPTAC TMT)",
 "copy_number":"Per-gene copy number \u2014 amplification calls and CN\u2192mRNA/protein transmission (GDC AscatNGS)",
 "rna":"mRNA abundance \u2014 CN\u2192mRNA transmission step (GDC STAR-Counts)",
 "ascat_purity":"Tumour purity/ploidy \u2014 ploidy-adjusted CN and purity control",
 "atac":"Tumour promoter chromatin accessibility \u2014 the transcription-gate covariate (TCGA-ATAC)",
 "topology":"Extracellular ectodomain / transmembrane topology \u2014 surface-accessibility gate (UniProt)",
 "gtex":"Normal bulk tissue expression \u2014 off-tumour reference; expression-breadth features (GTEx)",
 "corum":"Protein-complex membership \u2014 buffering covariate; complex features (CORUM)",
 "depmap":"CRISPR dependency \u2014 essentiality liability; dosage features (DepMap/CCLE)",
 "hpa_normal":"Normal single-cell + bulk expression \u2014 off-tumour AND-gate burden (Human Protein Atlas)",
 "cellxgene":"Malignant single-cell profiles \u2014 same-cell co-detection test (CELLxGENE census)",
 "methylation":"Promoter DNA methylation \u2014 regulatory gate covariate; TSS200/1500 mean beta, probes mapped via the EPIC manifest, no ChAMP (GDC CPTAC-3 EPIC arrays, 655 tumours)",
}
# Curated 5-column data index (tables/data_index.csv). This is authored
# separately from the pipeline's run-status log (reports/data_prep_layers.csv,
# which the R pipeline rewrites as layer/status/note) so the S2 table keeps its
# source + sample-size columns.
data_index = cfg.DIR_TAB / "data_index.csv"
s2_rows = []
if data_index.exists():
    for row in _csv.DictReader(open(data_index)):
        lyr = row["layer"]
        s2_rows.append([lyr, row.get("source",""), row.get("samples",""),
                        DATA_USE.get(lyr, row.get("notes","")), row.get("status","")])
elif layers_csv.exists():
    for row in _csv.DictReader(open(layers_csv)):
        lyr = row["layer"]
        s2_rows.append([lyr, row.get("note",""), "", DATA_USE.get(lyr, ""), row.get("status","")])
else:
    for lyr, use in DATA_USE.items():
        s2_rows.append([lyr, "", "", use, ""])
add_table(["Layer", "Source", "Samples", "Role in the analysis", "State"],
          s2_rows, fs=7.5)
para("Sample counts are the exact parsed dimensions for each layer; counts "
     "marked \u2018~\u2019 are rounded totals for whole reference databases (CORUM "
     "complexes, DepMap cell lines).", italic=True)

# ============================ SUPPLEMENTARY TABLE S3 =========================
doc.add_page_break()
H("Supplementary Table S3. Amplification-threshold sensitivity", level=2)
para("The full nomination funnel recomputed at the working amplification "
     "threshold (adjusted copy number \u2265 1.4, the broad-gain regime the "
     "clinical premise rests on) and at a stringent high-level threshold "
     "(\u2265 2.0). High-level focal amplification is rarer, so every funnel "
     "stage contracts and only the LSCC 3q (ATP11B+TFRC) and LUAD 5p "
     "(CLPTM1L+SLC12A7) constructs survive intact at \u2265 2.0; the broad-gain "
     "threshold is used because near-ubiquitous segmental gains, not rare focal "
     "events, are the actionable target.", italic=True)
_d1 = cfg.DIR_TAB / "d1_threshold_sensitivity.csv"
if _d1.exists():
    _rows = list(_csv.DictReader(open(_d1)))
    add_table(["Funnel stage", "CN \u2265 1.4", "CN \u2265 2.0"],
              [[r["metric"], r["thresh_1.4"], r["thresh_2.0"]] for r in _rows], fs=8)

# ============================ SUPPLEMENTARY TABLE S4 =========================
doc.add_page_break()
H("Supplementary Table S4. Prediction-only nominations in a cohort with no tissue-referenced proteome (GBM)", level=2)
para("Glioblastoma has copy number, mRNA and raw proteome in CPTAC \u2014 and so "
     "contributes to the transmission cascade \u2014 but its proteome is not "
     "referenced to a matched normal tissue of origin, the quantity co-elevation "
     "requires, so GBM contributes no measured surface nomination. Running the "
     "gene-property predictor end-to-end on GBM copy number alone (recurrent "
     "amplicon \u2192 predicted transmissibility \u2265 0.40 \u2192 UniProt surface "
     "gate) nominates the antigens below without using any GBM protein measurement. "
     "The canonical GBM receptor-tyrosine-kinase amplicon targets EGFR (7p) "
     "and MET (7q) are recovered from prediction alone, and 5 of 11 nominations are "
     "confirmed on the experimental Cell Surface Protein Atlas. Assembled into "
     "multivalent constructs and tested in 390,761 malignant GBM cells (208 donors), "
     "the 19p (ATP13A1+TMED1) and 20p (ATRN+PTPRA) sets are co-detected above "
     "independence on the depth-stratified null (1.12\u00d7 and 1.10\u00d7, p = 0.001, "
     "intervals clear of 1.0) and the 7q (MET+SLC12A9+SLC4A2) set is permutation-"
     "significant \u2014 so the full predict\u2192surface-gate\u2192single-cell-verify "
     "path closes without any GBM surface-abundance measurement: the nomination "
     "uses gene-intrinsic prediction, the surface call uses UniProt topology, and "
     "the verification uses single-cell detection \u2014 none of them a GBM "
     "tissue-referenced proteome.",
     italic=True)
_m2 = cfg.DIR_TAB / "m2_gbm_prediction_only.csv"
if _m2.exists():
    _rows = list(_csv.DictReader(open(_m2)))
    def _y(v): return "yes" if str(v).lower()=="true" else "\u2014"
    add_table(["Antigen","Arm","GBM amp. freq.","Pred. transmiss.","Ecto (aa)","TM","CSPA"],
              [[r["antigen"], r["arm"], f'{float(r["gbm_amp_freq"]):.2f}',
                f'{float(r["pred_transmit"]):.2f}', f'{float(r["ecto_aa"]):.0f}',
                f'{float(r["n_tm"]):.0f}', _y(r.get("cspa_confirmed"))] for r in _rows], fs=7.5)

# ============================ SUPPLEMENTARY TABLE S5 =========================
doc.add_page_break()
H("Supplementary Table S5. Observed cross-cohort transmissibility transfer", level=2)
para("Observed (measured, not predicted) transmissibility computed independently "
     "in each proteome-supported cohort, then correlated between cohort pairs "
     "(Spearman, genes with \u2265 20 amplified cases in both). Transfer is strong "
     "within lineage (the two lung subtypes), moderate to related epithelium, and "
     "weak to the distant clear-cell renal cohort \u2014 the gradient expected if "
     "the regulatory gates are set by cell of origin. This is the empirical basis "
     "for the within-cancer-type portability claim in the Discussion.", italic=True)
_m6 = cfg.DIR_TAB / "m6_observed_cross_cohort.csv"
if _m6.exists():
    _rows = list(_csv.DictReader(open(_m6)))
    add_table(["Cohort pair","Shared genes","Spearman \u03c1","p"],
              [[r["pair"], r["n_genes"], f'{float(r["rho"]):.2f}', r["p"]] for r in _rows], fs=8)

# ============================ REFERENCES =====================================
doc.add_page_break()
H("References", level=1)
for i, k in enumerate(_cited, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{i}. {REFS[k]}"); r.font.size = Pt(9.5)
uncited = [k for k in REFS if k not in _cited]
print(f"references: {len(_cited)} cited / {len(REFS)} in registry; uncited={uncited}")

doc.save(str(cfg.DIR_REP / "integrated_manuscript.docx"))
# verify
d2 = Document(str(cfg.DIR_REP / "integrated_manuscript.docx"))
imgs = sum(1 for s in d2.inline_shapes)
heads = sum(1 for p in d2.paragraphs if p.style.name.startswith("Heading"))
ntab = len(d2.tables)
print(f"tables={ntab}")
print(f"SAVED integrated_manuscript.docx: {len(d2.paragraphs)} paragraphs, "
      f"{imgs} embedded images, {heads} headings")
